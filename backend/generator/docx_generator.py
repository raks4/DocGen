from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


# ---------- INLINE FORMAT ----------
def apply_inline_format(paragraph, text):

    text = text.replace("** ", "**").replace(" **", "**")

    if text.count("**") % 2 != 0:
        text = text.replace("**", "")
    if text.count("*") % 2 != 0:
        text = text.replace("*", "")

    i = 0
    n = len(text)

    while i < n:

        # bold
        if text[i:i+2] == "**":
            end = text.find("**", i+2)
            if end != -1:
                run = paragraph.add_run(text[i+2:end])
                run.bold = True
                i = end + 2
                continue

        # italic
        if text[i] == "*" and (i+1 < n and text[i+1] != "*"):
            end = text.find("*", i+1)
            if end != -1:
                run = paragraph.add_run(text[i+1:end])
                run.italic = True
                i = end + 1
                continue

        # inline code
        if text[i] == "`":
            end = text.find("`", i+1)
            if end != -1:
                run = paragraph.add_run(text[i+1:end])
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(180, 0, 0)
                i = end + 1
                continue

        # links [text](url)
        if text[i] == "[":
            close = text.find("]", i)
            if close != -1 and text[close+1:close+2] == "(":
                end = text.find(")", close)
                if end != -1:
                    label = text[i+1:close]
                    run = paragraph.add_run(label)
                    run.font.color.rgb = RGBColor(0,102,204)
                    run.font.underline = True
                    i = end + 1
                    continue

        paragraph.add_run(text[i])
        i += 1


# ---------- CODE WRAP ----------
def wrap_code_lines(code, max_chars=90):
    wrapped = []

    for line in code.split("\n"):
        indent = len(line) - len(line.lstrip(" "))
        prefix = " " * indent
        text = line.lstrip()

        while len(text) > max_chars:
            part = text[:max_chars]
            wrapped.append(prefix + part)
            text = text[max_chars:]

        wrapped.append(prefix + text)

    return "\n".join(wrapped)


# ---------- CODE BLOCK ----------
def add_code_block(doc, code_text):

    code_text = wrap_code_lines(code_text, 90)

    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]

    # gray background
    tcPr = cell._element.xpath('./w:tcPr')[0]
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), "F7F7F8")
    tcPr.append(shd)

    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


# ---------- TABLE ----------
def add_table(doc, table_lines):

    rows=[]
    for row in table_lines:
        cols=[c.strip() for c in row.strip("|").split("|")]
        if all(set(c)<=set("-:") for c in cols):
            continue
        rows.append(cols)

    if not rows:
        return

    table=doc.add_table(rows=len(rows), cols=len(rows[0]))

    for i,r in enumerate(rows):
        for j,c in enumerate(r):
            p=table.rows[i].cells[j].paragraphs[0]
            apply_inline_format(p,c)


# ---------- DOCX CREATOR ----------
def create_docx(text):

    doc = Document()

    lines = [l for l in text.split("\n") if l.strip().lower() not in [
        "end of documentation",
        "end of document",
        "documentation ends here",
        "end."
    ]]

    in_code=False
    code_buffer=[]
    table_buffer=[]

    for line in lines:
        stripped=line.rstrip()

        if stripped.strip()=="---":
            continue

        # code block
        if stripped.strip().startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_buffer))
                code_buffer=[]
                in_code=False
            else:
                in_code=True
            continue

        if in_code:
            code_buffer.append(line)
            continue

        # table
        if "|" in stripped and stripped.count("|")>=2:
            table_buffer.append(stripped)
            continue
        else:
            if table_buffer:
                add_table(doc, table_buffer)
                table_buffer=[]

        # headings
        if stripped.startswith("# "):
            p=doc.add_paragraph(stripped[2:], style="Heading 1")
            for r in p.runs: r.font.color.rgb=RGBColor(11,61,145)
            continue

        if stripped.startswith("## "):
            p=doc.add_paragraph(stripped[3:], style="Heading 2")
            for r in p.runs: r.font.color.rgb=RGBColor(31,122,140)
            continue

        if stripped.startswith("### "):
            p=doc.add_paragraph(stripped[4:], style="Heading 3")
            for r in p.runs: r.font.color.rgb=RGBColor(138,90,68)
            continue

        # bullets
        if stripped.startswith("- "):
            p=doc.add_paragraph(style="List Bullet")
            apply_inline_format(p, stripped[2:])
            continue

        # paragraph
        if stripped.strip():
            p=doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            apply_inline_format(p, stripped)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
